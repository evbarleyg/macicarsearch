"""Excel + Word outputs for tools/build_site.py. Both take the computed model (same rows as the page)."""
import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter

# ---------------------------------------------------------------------------- xlsx
def _L(**k):
    k.setdefault('size', 10); return Font(name='Lora', **k)
BODY, BOLD, BLUE, MUTE = _L(), _L(bold=True), _L(color='0000FF'), _L(color='B7B7B7')
HDR, TITLE = _L(bold=True, color='FFFFFF'), Font(name='Lora', size=15, bold=True)
FILL_T = PatternFill('solid', start_color='B86046'); FILL_S = PatternFill('solid', start_color='D19B75')
FILL_I = PatternFill('solid', start_color='FFF2CC'); FILL_C = PatternFill('solid', start_color='F2E0BD'); FILL_G = PatternFill('solid', start_color='EFEFEF')
thin, med = Side('thin'), Side('medium')
BOX = Border(top=thin, bottom=thin, left=thin, right=thin)
CUR = '$#,##0_);($#,##0);"-"_)'; CUR2 = '$#,##0.00_);($#,##0.00);"-"_)'; INT = '#,##0_);(#,##0);"-"_)'; PCT2 = '0.00%_);(0.00%);"-"_)'
WRAP = Alignment(wrap_text=True, vertical='top')

def _sheet(wb, name, title, tab='B86046', first=False):
    ws = wb.active if first else wb.create_sheet()
    ws.title = name; ws.sheet_view.showGridLines = False; ws.sheet_properties.tabColor = tab
    ws.column_dimensions['A'].width = 3.5
    ws['B1'] = title; ws['B1'].font = TITLE
    for col in range(2, 16): ws.cell(row=1, column=col).border = Border(bottom=med)
    return ws

def _hdr(ws, row, labels, widths=None):
    for i, lab in enumerate(labels):
        c = ws.cell(row=row, column=2 + i, value=lab); c.font = HDR; c.fill = FILL_T; c.alignment = Alignment(vertical='center', wrap_text=True)
        if widths: ws.column_dimensions[get_column_letter(2 + i)].width = widths[i]
    ws.freeze_panes = ws.cell(row=row + 1, column=2)

def _row(ws, row, vals, fmts=None, font=BODY, wrap_cols=()):
    for i, v in enumerate(vals):
        c = ws.cell(row=row, column=2 + i, value=v); c.font = font
        if fmts and i < len(fmts) and fmts[i]: c.number_format = fmts[i]
        if i in wrap_cols: c.alignment = WRAP

def build_xlsx(m, path):
    B, META, H = m['B'], m['META'], m['helpers']
    money, d_short, tpl = H['money'], H['d_short'], H['tpl']
    wb = Workbook()

    # ---- Read me / control (first tab, black)
    ws = _sheet(wb, 'Read me', "CX-5 Seattle Buyer's Report — workbook", tab='000000', first=True)
    lines = [
        f'Published {H["d_long"](m["PUB"])}; refreshed {H["d_long"](m["REF"])}. Generated from data/board.json by tools/build_site.py — same rows as the web page.',
        f'{len(m["IN_PLAY"])} cars in play ({len(m["LIVE"])} live + benchmark), {len(m["SOLD"])} sold since publication, {len(m["QUOTED"])} written quotes in hand.',
        'Tabs: Board (live cars) · Quotes (every written quote, line items, counter, gap) · Monthly cost (blue inputs, live PMT formulas) · Sold-removed · Watchlist · Trim values · Sources.',
        'Font colours: blue = hard-coded input from a source (listing, dealer sheet, DOR rate); black = computed. Change the blue APR / term / down-payment drivers on Monthly cost and every payment recomputes.',
        f'Cheapest solid = {m["SOLID_RULE"]}. Stretch = {m["STRETCH_RULE"]}.',
        META['taxNote'],
        'Not affiliated with any dealer or listing site. Prices and availability change daily; verify history reports and the written OTD before purchase.',
    ]
    for i, t in enumerate(lines):
        c = ws.cell(row=3 + i, column=2, value=t); c.font = BODY; c.alignment = WRAP
    ws.column_dimensions['B'].width = 120
    r = 3 + len(lines) + 1
    ws.cell(row=r, column=2, value='What changed').font = BOLD
    for ch in B['changelog']:
        for b in ch['bullets']:
            r += 1; c = ws.cell(row=r, column=2, value=f'{H["d_long"](ch["date"])}: {b}'); c.font = BODY; c.alignment = WRAP

    # ---- Board
    ws = _sheet(wb, 'Board', f'Board — cars in play as of {H["d_long"](m["REF"])} (sold cars on their own tab)')
    labels = ['#', 'Status', 'Year', 'Trim', 'Engine', 'Miles', 'Listed $', 'Vehicle $ (pre-doc)', 'Doc fee', 'Dealer', 'City', 'Tax rate', 'Deal', 'KBB Seattle FPP', 'History', 'CPO', 'Days listed', 'Source', 'Listing URL', 'VIN', 'Stock', 'Comfort rank', 'Latest from dealer', 'Equipment']
    widths = [5, 10, 6, 28, 9, 9, 10, 12, 8, 26, 14, 8, 18, 12, 12, 6, 8, 11, 46, 20, 12, 8, 70, 60]
    _hdr(ws, 3, labels, widths)
    r = 4
    for c in m['IN_PLAY']:
        d = c['dealer']; dl = c['deal']
        vals = [c['rank'], c['status'], c['year'], c['trim'], c['engine'], c['miles'], c['price']['listed'], c['price']['advertised'], H['doc_fee'](c), d['name'], d['city'], H['tax_rate'](c) / 100,
                ' · '.join(x for x in [dl.get('rating'), dl.get('display')] if x), dl.get('kbbFpp'), H['hist_words'](c), 'Yes' if c['history'].get('cpo') else 'No', c.get('daysListed'), c['links']['primarySource'],
                c['links']['primary'], c.get('vin'), c.get('stock'), c.get('comfortRank'), tpl(c['latest']['text'], c), c.get('features')]
        _row(ws, r, vals, [INT, None, '0', None, None, INT, CUR, CUR, CUR, None, None, PCT2, None, CUR, None, None, INT, None, None, None, None, INT, None, None], wrap_cols=(22, 23))
        for col in (7, 8, 9, 12, 14): ws.cell(row=r, column=1 + col).font = BLUE
        r += 1

    # ---- Quotes
    ws = _sheet(wb, 'Quotes', 'Written quotes — line items as received, counter sent, gap')
    labels = ['#', 'Car', 'Dealer', 'Date', 'Document', 'Selling $', 'Doc $', 'Tax $', 'Tax rate', 'License / gov $', 'Taxes+fees lumped $', 'Add-ons', 'OTD (written)', 'First sheet OTD', 'Est. OTD at list', 'Counter OTD', 'Gap (written − counter)', 'Clean sheet?', 'Notes']
    widths = [5, 30, 26, 10, 44, 11, 8, 10, 8, 12, 14, 30, 13, 13, 13, 12, 14, 10, 70]
    _hdr(ws, 3, labels, widths)
    r = 4
    for c in [c for c in m['CARS'] if c.get('quote')]:
        q = c['quote']; cr = m['COST'].get(c['rank']); ctr = c.get('counter') or {}
        addons = ', '.join(a['name'] + (f' ${a["amount"]:,.0f}' if a.get('amount') else '') for a in (q.get('addOns') or [])) or 'none'
        vals = [c['rank'], f'{c["year"]} {c["trim"]}', c['dealer']['name'], d_short(q['date']) + (f' {q["time"]}' if q.get('time') else ''), q.get('docType'), q.get('selling'), q.get('doc'), q.get('tax'),
                (q['taxRate'] / 100) if q.get('taxRate') else None, q.get('license'), q.get('taxesAndFeesLumped'), addons, q.get('otd'), q.get('firstOtd'), cr['est'] if cr else None, ctr.get('otd'), None,
                'yes' if q.get('clean') else 'no', (q.get('note') or '') + (f' · counter: {ctr["note"]}' if ctr.get('note') else '')]
        _row(ws, r, vals, [INT, None, None, None, None, CUR2, CUR, CUR2, PCT2, CUR, CUR2, None, CUR2, CUR2, CUR, CUR, CUR, None, None], wrap_cols=(4, 11, 18))
        for col in (6, 7, 8, 9, 10, 11, 13, 14, 16): ws.cell(row=r, column=1 + col).font = BLUE
        if q.get('otd') and ctr.get('otd'):
            ws.cell(row=r, column=1 + 17, value=f'=N{r}-Q{r}').number_format = CUR
        r += 1

    # ---- Monthly cost (drivers + live formulas)
    ws = _sheet(wb, 'Monthly cost', 'Monthly cost — live cars, sorted by out-the-door (quoted where a dealer has replied)')
    ws['B3'] = 'APR'; ws['C3'] = m['APR'] / 100; ws['C3'].number_format = PCT2
    ws['B4'] = 'Term (months)'; ws['C4'] = m['N']
    ws['B5'] = 'Down payment (scenario 2)'; ws['C5'] = 5000; ws['C5'].number_format = CUR
    ws['B6'] = 'License / title / RTA est.'; ws['C6'] = META['licenseEst']; ws['C6'].number_format = CUR
    for rr in range(3, 7):
        ws.cell(row=rr, column=2).font = BOLD; c = ws.cell(row=rr, column=3); c.font = BLUE; c.fill = FILL_I; c.border = BOX
    ws['E3'] = 'Per $1,000 financed'; ws['E3'].font = BOLD; ws['F3'] = '=-PMT($C$3/12,$C$4,1000)'; ws['F3'].number_format = CUR2; ws['F3'].font = BODY
    ws['E4'] = f'{META["aprSource"]}; credit-union check {META["creditUnionCheck"]}'; ws['E4'].font = MUTE
    labels = ['#', 'Car', 'Dealer', 'Basis', 'Vehicle price', 'Tax rate', 'Doc fee', 'Est. OTD', 'Written OTD', 'OTD used', 'Monthly @ $0 down', 'Monthly @ down pmt', 'Total interest @ $0', 'Quote date / note']
    widths = [5, 30, 24, 8, 12, 8, 8, 12, 12, 12, 14, 14, 14, 40]
    _hdr(ws, 8, labels, widths)
    ws.freeze_panes = 'B9'
    r = 9
    for cr in m['COST_SORTED']:
        c = cr['car']
        _row(ws, r, [c['rank'], f'{c["year"]} {c["trim"]}' + (' (benchmark)' if c['status'] == 'benchmark' else ''), c['dealer']['name'], cr['basis'], cr['advertised'], cr['rate'] / 100, cr['doc'], None,
                     c['quote']['otd'] if cr['basis'] == 'quoted' else None, None, None, None, None,
                     (d_short(cr['date']) + ' sheet' if cr['basis'] == 'quoted' else f'est. at {c["dealer"]["city"]} rate') + (f' · if add-ons struck: {money(cr["struck"])}' if cr.get('struck') else '')],
             [INT, None, None, None, CUR, PCT2, CUR, CUR, CUR2, CUR2, CUR, CUR, CUR, None])
        for col in (5, 6, 7, 9): ws.cell(row=r, column=1 + col).font = BLUE
        ws.cell(row=r, column=1 + 8, value=f'=F{r}*(1+G{r})+H{r}+$C$6').number_format = CUR
        ws.cell(row=r, column=1 + 10, value=f'=IF(J{r}>0,J{r},I{r})').number_format = CUR2
        ws.cell(row=r, column=1 + 11, value=f'=-PMT($C$3/12,$C$4,K{r})').number_format = CUR
        ws.cell(row=r, column=1 + 12, value=f'=-PMT($C$3/12,$C$4,MAX(0,K{r}-$C$5))').number_format = CUR
        ws.cell(row=r, column=1 + 13, value=f'=L{r}*$C$4-K{r}').number_format = CUR
        for col in (8, 10, 11, 12, 13): ws.cell(row=r, column=1 + col).font = BODY
        if c['status'] == 'benchmark':
            for col in range(2, 16): ws.cell(row=r, column=col).fill = FILL_G
        r += 1
    r += 1
    ws.cell(row=r, column=2, value='Cheapest solid vs. stretch').font = BOLD
    cs, st = m['COST'][m['CHEAPEST']['rank']], m['COST'][m['STRETCH']['rank']]
    for lab, c, crr in (('Cheapest solid', m['CHEAPEST'], cs), ('Stretch', m['STRETCH'], st)):
        r += 1
        _row(ws, r, [None, lab, f'#{c["rank"]} {c["year"]} {c["trim"]} · {c["dealer"]["short"]}', crr['basis'], crr['price'], None, None, None, None, crr['otd'], crr['m0'], crr['m5'], crr['interest']],
             [None, None, None, None, CUR, None, None, None, None, CUR, CUR, CUR, CUR])
        ws.cell(row=r, column=3).font = BOLD
    r += 1
    _row(ws, r, [None, 'Δ', None, None, st['price'] - cs['price'], None, None, None, None, st['otd'] - cs['otd'], st['m0'] - cs['m0'], st['m5'] - cs['m5'], st['interest'] - cs['interest']],
         [None, None, None, None, CUR, None, None, None, None, CUR, CUR, CUR, CUR], font=BOLD)
    for col in range(2, 15): ws.cell(row=r, column=col).border = Border(top=thin, bottom=thin)
    r += 2
    c = ws.cell(row=r, column=2, value=f'Rule: cheapest solid = {m["SOLID_RULE"]}; stretch = {m["STRETCH_RULE"]}.'); c.font = MUTE

    # ---- Sold / removed
    ws = _sheet(wb, 'Sold-removed', 'Sold and removed since publication')
    _hdr(ws, 3, ['#', 'Year', 'Trim', 'Dealer', 'City', 'Was listed $', 'Miles', 'Sold', 'How confirmed', 'Listing URL'], [5, 6, 30, 30, 16, 12, 9, 10, 50, 50])
    r = 4
    for c in m['SOLD']:
        _row(ws, r, [c['rank'], c['year'], c['trimFull'], c['dealer']['name'], c['dealer']['city'], c['price']['listed'], c['miles'], d_short(c['soldInfo']['date']), c['soldInfo']['how'], c['links']['primary']],
             [INT, '0', None, None, None, CUR, INT, None, None, None])
        ws.cell(row=r, column=7).font = BLUE
        r += 1

    # ---- Watchlist
    ws = _sheet(wb, 'Watchlist', f'Screened and parked (re-checked {H["d_long"](m["REF"])})')
    _hdr(ws, 3, ['Car', 'Dealer', 'State', 'Price now', 'Was', 'Move', 'Days listed', 'Deal', 'Why parked', 'URL'], [42, 30, 6, 10, 10, 8, 8, 8, 60, 55])
    r = 4
    for w in B['watchlist']:
        l = w.get('latest') or {}
        _row(ws, r, [w['car'], w['dealer'], w['state'], l.get('price'), l.get('was'), l.get('move'), l.get('daysListed'), l.get('deal'), w['whyNot'], w['url']], [None, None, None, CUR, CUR, CUR, INT, None, None, None], wrap_cols=(8,))
        for col in (4, 5): ws.cell(row=r, column=1 + col).font = BLUE
        r += 1

    # ---- Trim values
    ws = _sheet(wb, 'Trim values', f'KBB fair purchase price vs. lowest clean live listing ({H["d_long"](m["REF"])})')
    _hdr(ws, 3, ['Year', 'Trim', "KBB nat'l FPP", 'KBB Seattle FPP', 'Lowest clean listing', 'Source', 'Δ vs Seattle (or nat’l)', 'vs national?', 'Listings'], [6, 24, 12, 14, 16, 12, 16, 10, 8])
    r = 4
    for t in m['TRIMS']:
        _row(ws, r, [t['year'], t['trim'], t['kbbNational'], t['kbbSeattle'], t['low'], t['lowSrc'], None, 'yes' if t['vsNat'] else '', t['n']], ['0', None, CUR, CUR, CUR, None, CUR, None, INT])
        for col in (3, 4, 5): ws.cell(row=r, column=1 + col).font = BLUE
        if t['low'] is not None:
            ws.cell(row=r, column=8, value=f'=F{r}-IF(E{r}>0,E{r},D{r})').number_format = CUR
        r += 1
    r += 1
    for f in B['prose']['trimFootnotes']:
        c = ws.cell(row=r, column=2, value=f); c.font = MUTE; r += 1

    # ---- Sources (last)
    ws = _sheet(wb, 'Sources', 'Sources & method')
    _hdr(ws, 3, ['Source', 'Status', 'Pulled (Aug 7)', 'What it added'], [16, 20, 12, 100])
    r = 4
    for s in B['sources']:
        _row(ws, r, [s['source'], s['statusLabel'], s['pulled'], s['added']], [None, None, INT, None], wrap_cols=(3,)); ws.cell(row=r, column=4).font = BLUE; r += 1
    r += 1
    ws.cell(row=r, column=2, value='Re-check sweeps').font = BOLD
    for s in B['sweeps']:
        r += 1; c = ws.cell(row=r, column=2, value=f'{H["d_long"](s["date"])}: {s["kind"]}. {s.get("method") or ""}'); c.font = BODY
    r += 2
    for t in [B['prose']['honestRead'], B['prose']['dedupe']] + B['prose']['filterRules'] + [B['prose']['recallNote']] + B['prose']['footnotes']:
        c = ws.cell(row=r, column=2, value=t); c.font = BODY; c.alignment = WRAP; r += 1
    ws.column_dimensions['B'].width = 16

    wb.move_sheet('Sources', offset=0)  # keep as last: move_sheet with 0 is a no-op; order is creation order
    wb.save(path)
    print(f'wrote {path.rsplit("/", 1)[-1]}')

# ---------------------------------------------------------------------------- docx
def build_docx(m, path):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    B, META, H = m['B'], m['META'], m['helpers']
    money, signed, d_short, d_long, tpl = H['money'], H['signed'], H['d_short'], H['d_long'], H['tpl']
    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10)

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
        for i, h in enumerate(headers): t.rows[0].cells[i].text = str(h)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = '' if v is None else str(v)
        doc.add_paragraph()
        return t

    doc.add_heading(META['title'], 0)
    doc.add_paragraph(META['scope'])
    p = doc.add_paragraph(f'Published {d_long(m["PUB"])} · refreshed {d_long(m["REF"])}: {len(m["IN_PLAY"])} cars in play, {len(m["SOLD"])} sold, {len(m["QUOTED"])} written quotes in hand. '
                          'This document, the web page, the status board and the Excel workbook are generated together from one board file; listings change daily, so verify before contacting.')

    doc.add_heading('TL;DR', 1)
    cs, stc = m['COST'][m['CHEAPEST']['rank']], m['COST'][m['STRETCH']['rank']]
    bullets = [
        f'Update {d_short(m["REF"])}: ' + B['changelog'][0]['bullets'][0],
        tpl(B['prose']['bestRemaining']),
        f'Cheapest solid vs. stretch: #{m["CHEAPEST"]["rank"]} {m["CHEAPEST"]["year"]} {m["CHEAPEST"]["trim"]} at {money(cs["otd"])} OTD ({cs["basis"]}, {money(cs["m0"])}/mo) vs. #{m["STRETCH"]["rank"]} {m["STRETCH"]["year"]} {m["STRETCH"]["trim"]} at {money(stc["otd"])} ({stc["basis"]}, {money(stc["m0"])}/mo).',
        f'Financing math: every $1,000 financed ≈ ${m["PER1K"]:.2f}/mo at {m["APR"]:.2f}% APR over {m["N"]} months. {META["cpoPromo"]}.',
    ]
    for b in bullets: doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('What changed', 1)
    for ch in B['changelog']:
        for b in ch['bullets']: doc.add_paragraph(f'{d_long(ch["date"])}: {b}', style='List Bullet')

    doc.add_heading(f'Top 3 right now ({d_short(m["REF"])})', 1)
    for c in m['PICKS'][:3]:
        cr = m['COST'][c['rank']]
        doc.add_heading(f'#{c["rank"]} · {c.get("pickTag") or ""} · {c["year"]} {c["trim"]} — {money(c["price"]["listed"])}', 2)
        doc.add_paragraph(f'{H["num"](c["miles"])} mi · {c["dealer"]["name"]} · {c["dealer"]["city"]} · {c["deal"].get("rating") or "—"} {c["deal"].get("display") or ""} · {H["hist_words"](c)}{" · CPO" if c["history"].get("cpo") else ""}')
        doc.add_paragraph(f'{money(cr["m0"])}/mo at $0 down on the {"written" if cr["basis"] == "quoted" else "estimated"} OTD of {money(cr["otd"])} ({m["APR"]:.2f}% · {m["N"]} mo).')
        doc.add_paragraph(tpl(c.get('pickWhy') or '', c))
        doc.add_paragraph(c['links']['primary'])
    doc.add_paragraph(tpl(B['prose']['runnerUp']))

    doc.add_heading(f'Shortlist and new candidates: {len(m["IN_PLAY"])} cars in play', 1)
    doc.add_paragraph(B['prose']['shortlistIntro'])
    rows = []
    for c in m['IN_PLAY']:
        qs = H['quoted_selling'](c)
        rows.append([f'#{c["rank"]}', c['year'], c['trim'], H['num'](c['miles']), (money(qs, True) + ' quoted; listed ' if qs else '') + H['price_listed_display'](c), f'{c["dealer"]["short"]}, {c["dealer"]["city"]}',
                     f'{c["deal"].get("rating") or "—"} {c["deal"].get("display") or ""}'.strip(), H['hist_short'](c), f'{c.get("daysListed")} d', tpl(c['latest']['text'], c)])
    table(['#', 'Year', 'Trim', 'Miles', 'Price', 'Seller', 'Deal', 'History', 'Listed', 'Latest'], rows)
    sold_line = 'Removed as sold: ' + ', '.join(f'#{c["rank"]} {c["dealer"]["short"]} {c["year"]} {c["trim"]} ({d_short(c["soldInfo"]["date"])})' for c in m['SOLD']) + '.'
    doc.add_paragraph(sold_line)
    doc.add_paragraph(B['prose']['notableDQ'])

    doc.add_heading('Which trim is the best buy', 1)
    doc.add_paragraph(B['prose']['trimIntro'])
    table(['Year', 'Trim', "KBB nat'l FPP", 'KBB Seattle FPP', f'Lowest clean listing ({d_short(m["REF"])})', 'Δ', 'Listings'],
          [[t['year'], t['trim'], money(t['kbbNational']), money(t['kbbSeattle']) if t['kbbSeattle'] else '—', (money(t['low']) + f' ({t["lowSrc"]})') if t['low'] else '—', (signed(t['delta']) + ('²' if t['vsNat'] else '')) if t['delta'] is not None else '—', t['n']] for t in m['TRIMS']])
    for f in B['prose']['trimFootnotes']: doc.add_paragraph(f)

    doc.add_heading('What it costs per month', 1)
    doc.add_paragraph(f'{m["N"]} months at {m["APR"]:.2f}% APR ({META["aprSource"]}); credit-union check {META["creditUnionCheck"]}. {META["taxNote"]}')
    table(['#', 'Car', 'Basis', 'Price', 'OTD', '@ $0 down', '@ $5k down', 'Interest @ $0'],
          [[f'#{r["car"]["rank"]}', f'{r["car"]["year"]} {r["car"]["trim"]} · {r["car"]["dealer"]["short"]}' + (' (benchmark)' if r['car']['status'] == 'benchmark' else ''),
            ('quoted ' + d_short(r['date'])) if r['basis'] == 'quoted' else f'est. {r["rate"]:.2f}%', money(r['price']), money(r['otd']) + (f' (if add-ons struck: {money(r["struck"])})' if r.get('struck') else ''), money(r['m0']), money(r['m5']), money(r['interest'])] for r in m['COST_SORTED']])
    doc.add_paragraph(f'Rule of thumb: every $1,000 financed ≈ ${m["PER1K"]:.2f}/mo at {m["APR"]:.2f}% / {m["N"]} mo.')
    doc.add_heading('Cheapest solid vs. stretch', 2)
    ch, sc = m['CHEAPEST'], m['STRETCH']
    table(['', 'Car', 'Price', 'OTD', '@ $0 down', '@ $5k down'],
          [['Cheapest solid', f'#{ch["rank"]} {ch["year"]} {ch["trim"]} · {ch["dealer"]["short"]}', money(cs['price']), f'{money(cs["otd"])} ({cs["basis"]})', money(cs['m0']), money(cs['m5'])],
           ['Stretch', f'#{sc["rank"]} {sc["year"]} {sc["trim"]} · {sc["dealer"]["short"]}', money(stc['price']), f'{money(stc["otd"])} ({stc["basis"]})', money(stc['m0']), money(stc['m5'])],
           ['Δ', '', signed(stc['price'] - cs['price']), signed(stc['otd'] - cs['otd']), signed(stc['m0'] - cs['m0']) + '/mo', signed(stc['m5'] - cs['m5']) + '/mo']])
    doc.add_paragraph(f'Cheapest solid = {m["SOLID_RULE"]}. Stretch = {m["STRETCH_RULE"]}.')
    d = m['CALC']
    doc.add_paragraph(f'Worked example, #{ch["rank"]} {ch["year"]} {ch["trim"]}: {money(d["price"])} price → {money(d["otd"])} out-the-door → {money(d["m0"])}/mo at $0 down, or {money(d["m5"])}/mo with $5,000 down ({m["APR"]:.2f}% APR, {m["N"]} mo).')

    doc.add_heading('Written quotes on file', 1)
    rows = []
    for c in [c for c in m['CARS'] if c.get('quote')]:
        q = c['quote']; ctr = c.get('counter') or {}
        rows.append([f'#{c["rank"]}', f'{c["year"]} {c["trim"]}', c['dealer']['name'], d_short(q['date']), money(q.get('selling'), True) if q.get('selling') else '—', money(q.get('otd'), True) if q.get('otd') else 'figures pending',
                     ', '.join(a['name'] for a in (q.get('addOns') or [])) or 'none', money(ctr.get('otd')) if ctr.get('otd') else '—', q.get('note') or ''])
    table(['#', 'Car', 'Dealer', 'Date', 'Selling', 'OTD', 'Add-ons', 'Counter', 'Notes'], rows)

    doc.add_heading('No-haggle ceiling: CarMax / Carvana', 1)
    doc.add_paragraph(B['prose']['noHaggleIntro'])
    for n in B['noHaggle']:
        doc.add_paragraph(f'{n["tier"]}: {n["seller"]} {money(n["price"]) if n.get("price") else ""} ({n["detail"]})' + (f' · Carvana {money(n["carvana"]["price"])} ({n["carvana"]["detail"]})' if n.get('carvana') else '') + (f'; {n["note"]}' if n.get('note') else ''), style='List Bullet')

    doc.add_heading('Email template: itemized OTD request', 1)
    doc.add_paragraph(B['prose']['email'])

    doc.add_heading('Sold and removed', 1)
    table(['#', 'Car', 'Dealer', 'Was listed', 'Sold'], [[f'#{c["rank"]}', f'{c["year"]} {c["trimFull"]}', c['dealer']['name'], money(c['price']['listed']), c['soldInfo']['how']] for c in m['SOLD']])

    doc.add_heading('Sources & method', 1)
    table(['Source', 'Status', 'Pulled', 'What it added'], [[s['source'], s['statusLabel'], s['pulled'], s['added']] for s in B['sources']])
    doc.add_paragraph(B['prose']['honestRead']); doc.add_paragraph(B['prose']['dedupe'])
    for s in B['sweeps']: doc.add_paragraph(f'{d_long(s["date"])}: {s["kind"]}. {s.get("method") or ""}', style='List Bullet')
    for r in B['prose']['filterRules']: doc.add_paragraph(r, style='List Bullet')
    doc.add_paragraph(B['prose']['recallNote'])
    for f in B['prose']['footnotes']: doc.add_paragraph(f, style='List Bullet')
    doc.add_paragraph(f'Research compiled {d_long(m["PUB"])} as a read-only scan; buyer outreach to dealers began Aug 11 and listings were last re-checked {d_long(m["REF"])}. Not affiliated with any seller. Verify history reports and pricing independently before purchase.')
    doc.save(path)
    print(f'wrote {path.rsplit("/", 1)[-1]}')
